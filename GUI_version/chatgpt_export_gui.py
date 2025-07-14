import tkinter as tk
from tkinter import filedialog, messagebox
from datetime import datetime, timezone
from pathlib import Path
import json
from docx import Document
import re
import os

def sanitize_filename(name):
    return re.sub(r'[\\/:*?"<>|]', '', name)

def extract_ordered_messages(mapping):
    root_id = next((k for k, v in mapping.items() if v.get("parent") is None), None)
    if not root_id:
        return []

    ordered = []
    current_id = mapping[root_id]["children"][0] if mapping[root_id]["children"] else None

    while current_id:
        node = mapping[current_id]
        message = node.get("message")
        if message:
            role = message["author"]["role"]
            parts = message["content"].get("parts", [])
            content = "\n".join(str(p) if isinstance(p, str) else json.dumps(p, ensure_ascii=False) for p in parts).strip()
            if content:
                ordered.append((role, content))
        children = node.get("children", [])
        current_id = children[0] if children else None

    return ordered

def save_as_markdown(messages, title, created, filepath):
    output = f"# {title}\n\nCreated: {created}\n\n"
    for role, content in messages:
        output += f"**{role.upper()}**:\n{content}\n\n"
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(output)

def save_as_docx(messages, title, created, filepath):
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Created: {created}")
    doc.add_paragraph("")
    for role, content in messages:
        doc.add_paragraph(f"{role.upper()}:", style='Heading3')
        doc.add_paragraph(content)
        doc.add_paragraph("")
    doc.save(filepath)

def process_conversations(input_path, output_dir):
    with open(input_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    conversations = data if isinstance(data, list) else data.get("conversations", [])

    Path(output_dir).mkdir(exist_ok=True)

    for convo in conversations:
        raw_title = convo.get("title", "Untitled").strip()
        title = sanitize_filename(raw_title.replace(" ", "_"))[:50]
        created = datetime.fromtimestamp(convo.get("create_time", 0), timezone.utc).strftime("%Y-%m-%d")
        filename_base = f"{created}_{title}"
        messages = extract_ordered_messages(convo["mapping"])
        if not messages:
            continue

        md_path = Path(output_dir) / f"{filename_base}.md"
        save_as_markdown(messages, raw_title, created, md_path)

        docx_path = Path(output_dir) / f"{filename_base}.docx"
        save_as_docx(messages, raw_title, created, docx_path)

    return len(conversations)

def run_gui():
    def select_input():
        path = filedialog.askopenfilename(filetypes=[("JSON files", "*.json")])
        if path:
            input_var.set(path)

    def select_output():
        path = filedialog.askdirectory()
        if path:
            output_var.set(path)

    def run():
        input_path = input_var.get()
        output_path = output_var.get()
        if not input_path or not output_path:
            messagebox.showerror("Missing Information", "Please select both input and output paths.")
            return
        try:
            count = process_conversations(input_path, output_path)
            messagebox.showinfo("Success", f"✅ Extracted {count} conversations into:\n{output_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Something went wrong:\n{str(e)}")

    root = tk.Tk()
    root.title("ChatGPT Export Reader")

    tk.Label(root, text="Select conversations.json:").grid(row=0, column=0, sticky="w", padx=5, pady=5)
    input_var = tk.StringVar()
    tk.Entry(root, textvariable=input_var, width=50).grid(row=0, column=1, padx=5)
    tk.Button(root, text="Browse", command=select_input).grid(row=0, column=2, padx=5)

    tk.Label(root, text="Select output folder:").grid(row=1, column=0, sticky="w", padx=5, pady=5)
    output_var = tk.StringVar()
    tk.Entry(root, textvariable=output_var, width=50).grid(row=1, column=1, padx=5)
    tk.Button(root, text="Browse", command=select_output).grid(row=1, column=2, padx=5)

    tk.Button(root, text="Convert", command=run, width=20, bg="green", fg="white").grid(row=2, column=1, pady=15)

    root.mainloop()

if __name__ == "__main__":
    run_gui()
